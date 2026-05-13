from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF

from services.llm_service import FunctionalRequirementDocument


def generate_docx(frd: FunctionalRequirementDocument) -> BytesIO:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Development Plan", 0)

    _add_section(doc, "Problem Summary", frd.problem_summary)
    _add_section(doc, "Proposed Solution", frd.proposed_solution)
    _add_section(doc, "Scope of Work", frd.scope_of_work)
    _add_section(doc, "User Flow", frd.user_flow)
    _add_section(doc, "Initial Architecture", frd.initial_architecture)

    # Feature Breakdown
    doc.add_heading("Feature Breakdown", level=1)
    modules: dict[str, list] = {}
    for feature in frd.feature_breakdown:
        modules.setdefault(feature.module, []).append(feature)
    for module, features in modules.items():
        doc.add_heading(module, level=2)
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light List Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text = "Feature"
        hdr[1].text = "Description"
        hdr[2].text = "Priority"
        for feature in features:
            row = table.add_row().cells
            row[0].text = feature.name
            row[1].text = feature.description
            row[2].text = feature.priority.replace("_", " ")
        doc.add_paragraph()

    # Timeline
    doc.add_heading("Timeline Estimation", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Phase"
    hdr[1].text = "Duration"
    hdr[2].text = "Deliverables"
    for phase in frd.timeline_estimation:
        row = table.add_row().cells
        row[0].text = phase.phase
        row[1].text = phase.duration
        row[2].text = "\n".join(f"- {d}" for d in phase.deliverables)

    # Risk Analysis
    doc.add_heading("Risk Analysis", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    hdr[0].text = "Risk"
    hdr[1].text = "Category"
    hdr[2].text = "Mitigation"
    for risk in frd.risk_analysis:
        row = table.add_row().cells
        row[0].text = risk.risk
        row[1].text = risk.category
        row[2].text = risk.mitigation

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def _add_section(doc: Document, title: str, text: str) -> None:
    doc.add_heading(title, level=1)
    for paragraph_text in text.split("\n"):
        paragraph_text = paragraph_text.strip()
        if paragraph_text:
            doc.add_paragraph(paragraph_text)


def generate_pdf(frd: FunctionalRequirementDocument) -> BytesIO:
    pdf = _PlanPDF()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 14, "Development Plan", ln=True, align="C")
    pdf.ln(6)

    _add_pdf_text_section(pdf, "Problem Summary", frd.problem_summary)
    _add_pdf_text_section(pdf, "Proposed Solution", frd.proposed_solution)
    _add_pdf_text_section(pdf, "Scope of Work", frd.scope_of_work)
    _add_pdf_text_section(pdf, "User Flow", frd.user_flow)
    _add_pdf_text_section(pdf, "Initial Architecture", frd.initial_architecture)

    # Feature Breakdown
    pdf.section_heading("Feature Breakdown")
    modules: dict[str, list] = {}
    for feature in frd.feature_breakdown:
        modules.setdefault(feature.module, []).append(feature)
    for module, features in modules.items():
        pdf.set_font("Helvetica", "B", 11)
        pdf.cell(0, 7, module, ln=True)
        pdf.set_font("Helvetica", "", 9)
        for feature in features:
            label = f"  {feature.name} [{feature.priority.replace('_', ' ')}]"
            pdf.cell(0, 5, label, ln=True)
            pdf.set_text_color(100, 100, 100)
            pdf.cell(0, 5, f"    {feature.description}", ln=True)
            pdf.set_text_color(0, 0, 0)
        pdf.ln(2)

    # Timeline
    pdf.section_heading("Timeline Estimation")
    col_widths = [65, 30, 95]
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(37, 99, 235)
    pdf.set_text_color(255, 255, 255)
    for text, w in zip(["Phase", "Duration", "Deliverables"], col_widths):
        pdf.cell(w, 7, text, 1, 0, "C", fill=True)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    for phase in frd.timeline_estimation:
        deliverables = ", ".join(phase.deliverables)
        max_lines = max(
            pdf.get_string_width(phase.phase) / (col_widths[0] - 2),
            pdf.get_string_width(phase.duration) / (col_widths[1] - 2),
            pdf.get_string_width(deliverables) / (col_widths[2] - 2),
            0,
        )
        row_h = max(7, int(max_lines * 5) + 7)
        pdf.cell(col_widths[0], row_h, phase.phase, 1)
        pdf.cell(col_widths[1], row_h, phase.duration, 1)
        pdf.cell(col_widths[2], row_h, deliverables, 1)
        pdf.ln()

    # Risk Analysis
    pdf.ln(4)
    pdf.section_heading("Risk Analysis")
    col_widths = [70, 25, 95]
    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(37, 99, 235)
    pdf.set_text_color(255, 255, 255)
    for text, w in zip(["Risk", "Category", "Mitigation"], col_widths):
        pdf.cell(w, 7, text, 1, 0, "C", fill=True)
    pdf.ln()
    pdf.set_font("Helvetica", "", 9)
    pdf.set_text_color(0, 0, 0)
    for risk in frd.risk_analysis:
        max_lines = max(
            pdf.get_string_width(risk.risk) / (col_widths[0] - 2),
            pdf.get_string_width(risk.category) / (col_widths[1] - 2),
            pdf.get_string_width(risk.mitigation) / (col_widths[2] - 2),
            0,
        )
        row_h = max(7, int(max_lines * 5) + 7)
        pdf.cell(col_widths[0], row_h, risk.risk, 1)
        pdf.cell(col_widths[1], row_h, risk.category, 1)
        pdf.cell(col_widths[2], row_h, risk.mitigation, 1)
        pdf.ln()

    buf = BytesIO()
    pdf.output(buf)
    buf.seek(0)
    return buf


def _add_pdf_text_section(pdf: FPDF, title: str, text: str) -> None:
    pdf.section_heading(title)
    pdf.set_font("Helvetica", "", 10)
    pdf.multi_cell(0, 5, text)
    pdf.ln(3)


class _PlanPDF(FPDF):
    def section_heading(self, title: str) -> None:
        self.set_font("Helvetica", "B", 13)
        self.set_text_color(37, 99, 235)
        self.cell(0, 10, title, ln=True)
        self.set_text_color(0, 0, 0)
